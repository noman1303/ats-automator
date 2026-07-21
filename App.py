"""
=====================================================================
 ADVANCED ATS RESUME AUTOMATOR  v4 — BATCH MODE + COVER LETTERS
 (Streamlit + OpenAI-compatible APIs + python-docx)
=====================================================================
 Developed by Noman Belim
=====================================================================
 WHAT'S NEW IN v4
 - BATCH MODE: process many job descriptions in a single run instead
   of one at a time. Paste several JDs separated by a marker line, or
   upload one .txt file per JD. The app tailors the resume for every
   JD, then hands back ONE zip containing one company/ subfolder per
   application (built for people applying to a lot of jobs per day).
   A failure on one JD (bad key, malformed JD, etc.) is caught and
   logged per-row instead of stopping the whole batch, and a
   summary.csv is included in the zip so you can see coverage % and
   status for every application at a glance.
 - AUTO COVER LETTERS: an optional checkbox generates a tailored,
   non-fabricated cover letter alongside every resume, in the SAME
   AI call as the resume tailoring (no extra API call, so it doesn't
   cost you extra quota at volume). Works in both single and batch
   mode. Cover letters use only real facts from the candidate's
   profile — same "never fabricate" rule as the resume.
 - Everything from v3 (multi-key Gemini pool, multi-provider
   failover, 8-stage tailoring, locked facts, code-level keyword
   coverage) is unchanged.

 WHAT'S NEW IN v3
 - MULTI-KEY GEMINI POOL: paste ALL your Gemini API keys (one per
   line) into a text file (e.g. gemini_keys.txt). The app loads them
   automatically and rotates through them — when one key gets
   rate-limited, it's put on cooldown and the app instantly tries
   the next key in your file. No more copy-pasting keys by hand.
 - Still supports the other 4 providers (Cerebras, Groq, Mistral,
   OpenRouter) as single-key fallbacks after the Gemini pool is
   exhausted.
 - Everything else (8-stage tailoring, locked facts, ATS docx,
   output/ folder, code-level keyword coverage) is
   unchanged from v2.

 HOW TO USE THE KEY FILE:
   1. Create a plain .txt file, e.g. C:\noman\gemini_keys.txt
   2. Paste one Gemini API key per line (blank lines / lines
      starting with # are ignored):

        AIzaSyABC123...
        AIzaSyDEF456...
        # this one is my backup account
        AIzaSyGHI789...

   3. In the sidebar, point "Gemini keys file" at that path.
   4. Whenever a key gets rate-limited, the app skips it for 15
      minutes and moves to the next key automatically — you never
      touch the app again until all keys in the file are cooling
      down.
   5. Add more keys any time by editing the .txt file and clicking
      "🔄 Reload keys from file" in the sidebar (no restart needed).

 =====================================================================
 HOW TO RUN THIS APP (Windows) — copy-paste these, in order
 =====================================================================
 1) Open Command Prompt.

 2) Go to the folder where this file lives:
        cd D:\\Noman

 3) (One-time / whenever packages need installing or updating)
    If plain "pip" isn't recognized, use "python -m pip" instead:
        python -m pip install streamlit openai python-docx

 4) Run the app.
    If plain "streamlit" isn't recognized (common when Python isn't
    fully added to PATH), run it as a module instead — this always
    works as long as "python" itself works:
        python -m streamlit run app.py

 5) It opens in your browser automatically at something like:
        http://localhost:8501

 6) To stop the app: click back into the Command Prompt window and
    press Ctrl + C.

 TROUBLESHOOTING QUICK REFERENCE:
   - "'pip' is not recognized"       -> use: python -m pip install ...
   - "'streamlit' is not recognized" -> use: python -m streamlit run app.py
   - Check Python is installed at all -> python --version
=====================================================================
"""

import os
import re
import csv
import json
import time
import tempfile
import io
import zipfile
import base64
from pathlib import Path
from datetime import datetime

import streamlit as st
from openai import OpenAI
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT

# ---------------------------------------------------------------
# APP METADATA / CREDIT
# ---------------------------------------------------------------
APP_AUTHOR = "Noman Belim"
APP_VERSION = "v4"

# ---------------------------------------------------------------
# BATCH MODE — line used to separate multiple pasted JDs in one
# text box. Any JD block may optionally start with a line like
# "COMPANY: Acme Corp" to force the output folder/company name,
# for JDs (e.g. from a recruiter) that don't clearly state the
# hiring company in the JD text itself.
# ---------------------------------------------------------------
JD_SPLIT_MARKER = "===NEXT JD==="
FORCE_COMPANY_PREFIX = "company:"

# ---------------------------------------------------------------
# CANDIDATE PATH PICKER — add your own saved candidates here.
# Dropdown shows these first; "Custom / one-off path" falls back
# to a free-text field for anything not in this list.
# ---------------------------------------------------------------
SAVED_CANDIDATES = {
    "Bethlehem Lulseged": r"D:\Noman\BETHLEHEM\Bethlehem_Lulseged_Resume.pdf",
    # "Another Candidate": r"D:\Noman\OTHER\Another_Candidate_Resume.docx",
}
CUSTOM_PATH_LABEL = "— Custom / one-off path —"

# ---------------------------------------------------------------
# PROVIDER CHAIN — tried top to bottom. Reorder to change priority.
# Gemini is special-cased below to support a POOL of keys from a
# text file; the other providers still use a single key each from
# the sidebar, exactly like v2.
# ---------------------------------------------------------------
PROVIDERS = [
    {
        "id": "gemini", "label": "Google Gemini",
        "base_url": "https://generativelanguage.googleapis.com/v1beta/openai/",
        "models": ["gemini-3-flash-preview", "gemini-3.1-flash-lite", "gemini-2.5-flash"],
        "hint": "Free key: aistudio.google.com → Get API key",
    },
    
]

GEMINI_BASE_URL = PROVIDERS[0]["base_url"]
GEMINI_MODELS = PROVIDERS[0]["models"]

COOLDOWN_SECONDS = 15 * 60   # skip a rate-limited provider/key for 15 min
RATE_LIMIT_MARKERS = ("429", "rate", "quota", "exceed", "resource_exhausted",
                      "capacity", "overloaded", "limit")
MODEL_GONE_MARKERS = ("404", "not found", "no longer available", "deprecated",
                      "decommissioned", "does not exist", "invalid model")

st.set_page_config(page_title="ATS Resume Automator", page_icon="⚡", layout="wide")

# ---------------------------------------------------------------
# GLOBAL THEME — "Aurora Violet"
# Dark canvas + violet→pink gradient accents (matches the credit
# badge palette). Paired with .streamlit/config.toml for native
# widget colors (buttons, sliders, checkboxes, focus rings).
# ---------------------------------------------------------------
st.markdown(
    """
    <style>
    /* ---------- page canvas ---------- */
    .stApp {
        background:
            radial-gradient(circle at 15% -10%, rgba(139,92,246,0.16), transparent 45%),
            radial-gradient(circle at 85% 110%, rgba(236,72,153,0.14), transparent 45%),
            #0e0e17;
    }

    /* ---------- sidebar ---------- */
    section[data-testid="stSidebar"] {
        background: linear-gradient(180deg, #14141f 0%, #191927 100%);
        border-right: 1px solid rgba(139,92,246,0.18);
    }

    /* ---------- headings ---------- */
    h1 {
        background: linear-gradient(135deg, #8b5cf6 0%, #6366f1 40%, #ec4899 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        background-clip: text;
        font-weight: 800 !important;
    }
    h2 {
        color: #c4b5fd !important;
        font-weight: 700 !important;
        border-left: 4px solid #8b5cf6;
        padding-left: 10px;
    }
    h3 {
        color: #e9d5ff !important;
        font-weight: 650 !important;
    }

    /* ---------- buttons ---------- */
    .stButton > button, .stDownloadButton > button {
        border-radius: 10px !important;
        font-weight: 600 !important;
        border: 1px solid rgba(139,92,246,0.35) !important;
        transition: all 0.18s ease !important;
    }
    .stButton > button:hover, .stDownloadButton > button:hover {
        transform: translateY(-1px);
        box-shadow: 0 6px 18px rgba(139,92,246,0.35);
        border-color: #ec4899 !important;
    }
    button[kind="primary"] {
        background: linear-gradient(135deg, #8b5cf6 0%, #6366f1 45%, #ec4899 100%) !important;
        box-shadow: 0 3px 14px rgba(139,92,246,0.4) !important;
    }

    /* ---------- inputs ---------- */
    .stTextInput input, .stTextArea textarea {
        border-radius: 8px !important;
        border: 1px solid rgba(139,92,246,0.3) !important;
    }
    .stTextInput input:focus, .stTextArea textarea:focus {
        border-color: #ec4899 !important;
        box-shadow: 0 0 0 2px rgba(236,72,153,0.22) !important;
    }

    /* ---------- expanders ---------- */
    [data-testid="stExpander"] {
        border: 1px solid rgba(139,92,246,0.22) !important;
        border-radius: 12px !important;
        background: rgba(139,92,246,0.05) !important;
        overflow: hidden;
    }

    /* ---------- metrics ---------- */
    [data-testid="stMetric"] {
        background: linear-gradient(135deg, rgba(139,92,246,0.12), rgba(236,72,153,0.10));
        border: 1px solid rgba(139,92,246,0.3);
        border-radius: 14px;
        padding: 14px 16px;
    }

    /* ---------- alerts (success / error / warning / info) ---------- */
    [data-testid="stAlert"] {
        border-radius: 10px !important;
    }

    /* ---------- dividers ---------- */
    hr {
        border: none !important;
        height: 1px !important;
        background: linear-gradient(90deg, transparent, rgba(139,92,246,0.5), transparent) !important;
        margin: 1.3rem 0 !important;
    }

    /* ---------- status widget ---------- */
    [data-testid="stStatusWidget"] {
        border-radius: 12px !important;
        border: 1px solid rgba(139,92,246,0.25) !important;
    }

    /* ---------- scrollbar ---------- */
    ::-webkit-scrollbar { width: 10px; height: 10px; }
    ::-webkit-scrollbar-track { background: #14141f; }
    ::-webkit-scrollbar-thumb {
        background: linear-gradient(180deg, #8b5cf6, #ec4899);
        border-radius: 10px;
    }
    </style>
    """,
    unsafe_allow_html=True,
)


# ---------------------------------------------------------------
# 0. GEMINI KEY LOADER
#    Reads one API key per line from either an uploaded .txt file
#    (works everywhere, including Streamlit Cloud) or a local file
#    path (only works when the app runs on your own computer).
#    Blank lines and lines starting with # are ignored so you can
#    leave yourself notes next to each key.
# ---------------------------------------------------------------
def parse_keys_text(text: str) -> list[str]:
    """Return a de-duplicated, ordered list of keys from raw text (one per line)."""
    keys, seen = [], set()
    for line in text.splitlines():
        key = line.strip()
        if not key or key.startswith("#"):
            continue
        if key not in seen:
            seen.add(key)
            keys.append(key)
    return keys


def load_keys_from_file(file_path: str) -> list[str]:
    """
    Return a de-duplicated, ordered list of keys from a LOCAL text file.
    NOTE: this only works when Streamlit is running on the same machine
    where the file lives. On Streamlit Cloud (or any hosted deployment)
    the app runs on a remote server that has no access to your PC's
    C:\\ drive or to a URL you paste in — use the file uploader instead.
    """
    p = Path(file_path.strip().strip('"'))
    if not p.exists():
        raise FileNotFoundError(f"Key file not found: {p}")
    return parse_keys_text(p.read_text(encoding="utf-8", errors="ignore"))


def gemini_key_short(key: str) -> str:
    """Short label for logs/UI so full keys are never printed on screen."""
    if len(key) <= 10:
        return key[:4] + "…"
    return f"{key[:6]}…{key[-4:]}"


# ---------------------------------------------------------------
# 1. TEXT EXTRACTION  (reads the original resume ONE time)
# ---------------------------------------------------------------
def extract_text_from_file(file_path: Path) -> str:
    """Read raw text from a .docx or .pdf resume."""
    suffix = file_path.suffix.lower()
    if suffix == ".docx":
        doc = Document(str(file_path))
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        return "\n".join(t for t in parts if t.strip())
    elif suffix == ".pdf":
        from pypdf import PdfReader
        reader = PdfReader(str(file_path))
        return "\n".join((page.extract_text() or "") for page in reader.pages)
    else:
        raise ValueError("Unsupported file type. Please use .docx or .pdf")


# ---------------------------------------------------------------
# 2. MULTI-PROVIDER + MULTI-KEY AI CALL WITH AUTOMATIC FAILOVER
#    Order of attempts:
#      1) Every Gemini key from the key-file pool (round-robin,
#         skipping keys on cooldown), each tried against every
#         Gemini model in turn.
#      2) A single sidebar-entered Gemini key, if the user also
#         typed one directly (kept for backward compatibility).
#      3) The remaining single-key providers (Cerebras, Groq,
#         Mistral, OpenRouter), same as v2.
# ---------------------------------------------------------------
def extract_json(raw: str) -> dict:
    """Pull a JSON object out of a model reply, tolerating fences/extra text."""
    raw = re.sub(r"^```(json)?|```$", "", raw.strip(), flags=re.MULTILINE).strip()
    start, end = raw.find("{"), raw.rfind("}")
    if start == -1 or end == -1:
        raise ValueError("No JSON object found in AI reply")
    return json.loads(raw[start:end + 1])


def _try_gemini_key(key: str, prompt: str, cooldowns: dict, notes: list, errors: list):
    """
    Try one Gemini key across all Gemini models.
    Returns (data, provider_label) on success, or None on failure
    (cooldowns/notes/errors are mutated in place).
    """
    now = time.time()
    short = gemini_key_short(key)
    wait = cooldowns.get(f"gemini:{key}", 0) - now
    if wait > 0:
        notes.append(f"Gemini key {short}: on cooldown ({int(wait // 60) + 1} min left)")
        return None

    client = OpenAI(api_key=key, base_url=GEMINI_BASE_URL, timeout=90)
    for model in GEMINI_MODELS:
        try:
            resp = client.chat.completions.create(
                model=model,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.3,
                max_tokens=4096,
            )
            data = extract_json(resp.choices[0].message.content)
            return data, f"Gemini key {short} ({model})"
        except Exception as e:                        # noqa: BLE001
            msg = str(e)
            errors.append(f"Gemini {short} / {model}: {msg[:140]}")
            low = msg.lower()
            if any(m in low for m in MODEL_GONE_MARKERS):
                notes.append(f"Gemini key {short}: model '{model}' unavailable → trying next model")
                continue
            if any(m in low for m in RATE_LIMIT_MARKERS):
                cooldowns[f"gemini:{key}"] = time.time() + COOLDOWN_SECONDS
                notes.append(f"Gemini key {short}: rate-limited → switching to next key")
                return None
            notes.append(f"Gemini key {short} / {model}: error → trying next option")
            continue
    return None


def call_ai_json(keys: dict, gemini_key_pool: list, prompt: str):
    """
    Try, in order:
      1) every key in gemini_key_pool (skipping ones on cooldown)
      2) a manually-entered Gemini key from the sidebar (if any)
      3) the other single-key providers
    Returns (parsed_json, provider_label, notes).
    """
    cooldowns = st.session_state.setdefault("cooldowns", {})
    notes, errors = [], []

    # --- 1) Rotate through the Gemini key-file pool -------------
    for key in gemini_key_pool:
        result = _try_gemini_key(key, prompt, cooldowns, notes, errors)
        if result is not None:
            data, label = result
            return data, label, notes

    # --- 2) Fall back to a manually-typed Gemini key, if present -
    manual_gemini_key = (keys.get("gemini") or "").strip()
    if manual_gemini_key and manual_gemini_key not in gemini_key_pool:
        result = _try_gemini_key(manual_gemini_key, prompt, cooldowns, notes, errors)
        if result is not None:
            data, label = result
            return data, label, notes

    # --- 3) Fall back to the other single-key providers ----------
    now = time.time()
    for p in PROVIDERS:
        if p["id"] == "gemini":
            continue  # already handled above via the key pool / manual key
        key = (keys.get(p["id"]) or "").strip()
        if not key:
            continue
        wait = cooldowns.get(p["id"], 0) - now
        if wait > 0:
            notes.append(f"{p['label']}: on cooldown ({int(wait // 60) + 1} min left)")
            continue
        client = OpenAI(api_key=key, base_url=p["base_url"], timeout=90)
        rate_limited = False
        for model in p["models"]:
            try:
                resp = client.chat.completions.create(
                    model=model,
                    messages=[{"role": "user", "content": prompt}],
                    temperature=0.3,
                    max_tokens=4096,
                )
                data = extract_json(resp.choices[0].message.content)
                return data, f"{p['label']} ({model})", notes
            except Exception as e:                    # noqa: BLE001
                msg = str(e)
                errors.append(f"{p['label']} / {model}: {msg[:140]}")
                low = msg.lower()
                if any(m in low for m in MODEL_GONE_MARKERS):
                    notes.append(f"{p['label']}: model '{model}' unavailable → trying next model")
                    continue
                if any(m in low for m in RATE_LIMIT_MARKERS):
                    cooldowns[p["id"]] = time.time() + COOLDOWN_SECONDS
                    notes.append(f"{p['label']}: rate-limited → switching to next provider")
                    rate_limited = True
                    break
                notes.append(f"{p['label']} / {model}: error → trying next option")
                continue
        if rate_limited:
            continue

    raise RuntimeError(
        "All configured AI providers/keys failed or are on cooldown.\n\n" +
        "\n".join(errors + notes) +
        "\n\nFixes: wait a few minutes, add more Gemini keys to your key file, "
        "add another provider key in the sidebar, or check your internet connection."
    )


# ---------------------------------------------------------------
# 3. STEP ONE — PARSE ORIGINAL RESUME INTO A STRUCTURED PROFILE
#    (cached to _profile.json → only 1 AI call per candidate, ever)
# ---------------------------------------------------------------
PARSE_PROMPT = """You are a resume parsing engine. Convert the resume text below into strict JSON.
Return ONLY JSON, no commentary. Use this exact schema:

{{
  "name": "",
  "headline": "",
  "contact_line": "",
  "summary": "",
  "skills": [ {{"category": "", "items": ""}} ],
  "experience": [
    {{
      "title": "",
      "company": "",
      "dates": "",
      "bullets": ["", ""]
    }}
  ],
  "education": ["", ""]
}}

Rules:
- Copy job titles, company names, dates, education, and certifications EXACTLY as written.
- Keep every bullet point of every job.
- "contact_line" = location | phone | email | linkedin joined with " | ".
- "education" = one string per degree/certification line.

RESUME TEXT:
----------------
{resume_text}
----------------
"""


def profile_cache_path(resume_path: Path) -> Path:
    return resume_path.parent / (resume_path.stem + "_profile.json")


def load_or_parse_profile(keys: dict, gemini_key_pool: list, resume_path: Path, force: bool = False):
    cache = profile_cache_path(resume_path)
    if cache.exists() and not force:
        return json.loads(cache.read_text(encoding="utf-8")), None
    text = extract_text_from_file(resume_path)
    profile, provider, _ = call_ai_json(keys, gemini_key_pool, PARSE_PROMPT.format(resume_text=text))
    cache.write_text(json.dumps(profile, indent=2), encoding="utf-8")
    return profile, provider


# ---------------------------------------------------------------
# 4. STEP TWO — TAILOR THE PROFILE TO A SPECIFIC JD (8-STAGE PROCESS)
# ---------------------------------------------------------------
TAILOR_PROMPT = """You are an elite ATS resume optimization engine used by a professional recruiter.
Tailor the candidate's resume to the Job Description by working through this 8-STAGE PROCESS, in order.

STAGE 1 — DECONSTRUCT THE JD FIRST:
Extract five things: (a) the exact job title, (b) hard skills & tools, (c) process language
(e.g. backlog, risk/issue logs, escalation, status reporting), (d) the domain/industry,
(e) soft requirements (e.g. matrixed environment, self-motivated). Rank keywords by priority:
anything in the job title, in the first paragraph, or repeated 2+ times is TOP PRIORITY and
must appear in the summary, the skills, AND at least one experience bullet.

STAGE 2 — LOCKED vs EDITABLE ZONES:
LOCKED (copy through completely unchanged): name, headline, contact_line, job titles,
company names, dates, education, skill category names and their order, the number of bullets
per job, and every real metric (%, $, team sizes) — numbers are what recruiters trust.
EDITABLE: summary wording, skill items inside existing categories, bullet wording.

STAGE 3 — GAP ANALYSIS (the most important stage). Sort every top JD requirement into 3 buckets:
- DIRECT MATCH: the candidate already does it -> rewrite it using the JD's exact phrase
  (e.g. resume says "risk logs", JD says "risk/issue logs" -> use "risk/issue logs").
- IMPLIED MATCH: the candidate clearly did it but never used the JD's words -> make the
  implied work explicit (e.g. "facilitated stand-ups and retrospectives" implies
  "documented meeting notes, decisions, and follow-up actions").
- NO MATCH: the candidate does not have it -> DO NOT FABRICATE. Never invent employers,
  clients, domain experience (e.g. Medicaid), certifications, degrees, or tools.
  Only emphasize genuinely adjacent real experience, then stop.

STAGE 4 — SUMMARY FORMULA (4-5 sentences):
[Role aligned to JD title] + [years of experience] + [real scale numbers] +
[top 5-6 JD keywords woven in naturally] + [outcome language].
The summary carries the heaviest keyword load.

STAGE 5 — SKILLS REBUILD:
Same categories, same order. Inside each category put JD-priority terms first. Include both
forms of a term where useful — "Microsoft Project (MS Project)", acronym plus spelled-out
"Risk/Issue Logs (RAID)" — because different ATS systems match differently. Only add a skill
if it logically follows from real experience (someone running Jira sprints can honestly
list "Backlog Management").

STAGE 6 — BULLET FORMULA (keep the same bullet count per job):
Each bullet = strong action verb + task in JD terminology + tool + real metric or outcome.
Maximum 1-2 JD keywords per bullet (more = stuffing, and modern ATS penalizes it).
Never open two bullets in the same job with the same verb.
Spread the top keywords across ALL jobs so density looks natural, not dumped in one place.

STAGE 7 — ATS COMPLIANCE CHECK:
Target roughly 70-80% coverage of the top JD keywords across the whole document.
100% coverage is a red flag to recruiters, not a goal. Honestly compute your coverage.

STAGE 8 — HUMAN READ-THROUGH:
Read the result as a recruiter, not a machine. Every line must sound human-written and be
defensible by the candidate in an interview. If a line fails either test, soften it.
{cover_letter_stage}
OUTPUT: Return ONLY a JSON object — the FULL updated profile in the SAME schema as the
input profile, PLUS these extra top-level keys:
  "job_title_detected": short job title from the JD (used for the file name),
  "company_detected": hiring company from the JD, or "Company" if unknown,
  "matched_keywords": top 15 JD keywords now present in the resume,
  "missing_keywords": top JD requirements deliberately NOT added (the NO MATCH bucket),
  "gap_analysis": {{
      "direct_matches": ["..."],
      "implied_matches": ["..."],
      "no_match_not_fabricated": ["..."]
  }}{cover_letter_key}

Note: keyword coverage percentage is calculated separately in code after you respond —
you do not need to compute or report it yourself.

(A) CANDIDATE PROFILE JSON:
{profile_json}

(B) JOB DESCRIPTION:
----------------
{jd_text}
----------------
"""

# Only injected into the prompt when the user checks "generate cover letter" —
# keeps the prompt (and token cost) smaller for people who don't want one.
COVER_LETTER_STAGE = """
STAGE 9 — COVER LETTER (requested for this run):
Write one complete, ready-to-send cover letter for THIS specific role at THIS specific
company, using ONLY real facts from the candidate profile above (same locked-facts rule
as the resume — never invent an employer, metric, or achievement).
- 3-4 short paragraphs, roughly 250-350 words total, in plain prose (no bullet points).
- Paragraph 1: name the role and company, one line on who the candidate is (title + years).
- Paragraphs 2-3: 2-3 of the strongest DIRECT MATCH / IMPLIED MATCH points from Stage 3,
  written as short stories with real metrics — not a restatement of the resume bullets.
- Final paragraph: brief, confident close + availability + thanks.
- Salutation "Dear Hiring Manager," unless the JD clearly names a person.
- Return it as a single string in the "cover_letter" key, with "\\n\\n" between paragraphs.
"""


def tailor_profile(keys: dict, gemini_key_pool: list, profile: dict, jd_text: str,
                    include_cover_letter: bool = False):
    cover_letter_key = (
        '\n  "cover_letter": full tailored cover letter text per the cover-letter stage '
        'below, "\\n\\n"-separated paragraphs.'
        if include_cover_letter else ""
    )
    prompt = TAILOR_PROMPT.format(
        cover_letter_stage=COVER_LETTER_STAGE if include_cover_letter else "",
        cover_letter_key=cover_letter_key,
        profile_json=json.dumps(profile, indent=1),
        jd_text=jd_text[:15000],
    )
    return call_ai_json(keys, gemini_key_pool, prompt)


# ---------------------------------------------------------------
# 4b. CODE-LEVEL KEYWORD COVERAGE  (replaces AI self-reported %)
#     Extracts the highest-frequency meaningful words/phrases from
#     the JD, then checks how many actually appear in the tailored
#     resume text. This is deterministic and can't be "faked" by
#     the model the way a self-reported percentage could be.
# ---------------------------------------------------------------
STOPWORDS = {
    "the", "and", "a", "an", "to", "of", "in", "for", "on", "with", "as", "is",
    "at", "by", "or", "be", "this", "that", "will", "are", "you", "your", "our",
    "we", "from", "have", "has", "it", "its", "into", "such", "who", "may",
    "can", "all", "their", "these", "those", "if", "not", "than", "then",
    "them", "they", "he", "she", "his", "her", "which", "about", "including",
    "etc", "per", "up", "out", "over", "under", "job", "role", "position",
    "candidate", "candidates", "applicant", "applicants", "company", "team",
    "years", "year", "experience", "ability", "strong", "must", "required",
    "requirements", "preferred", "skills", "work", "working", "including",
}


def extract_jd_keywords(jd_text: str, top_n: int = 25) -> list[str]:
    """
    Pull the top_n highest-frequency, meaningful words/short-phrases
    from the JD text. Combines single words and 2-word phrases so
    multi-word terms (e.g. "risk log", "stand up") aren't lost.
    """
    text = jd_text.lower()
    words = re.findall(r"[a-z][a-z0-9+#./-]*", text)
    words = [w.strip(".-/") for w in words if w.strip(".-/")]

    freq = {}
    for w in words:
        if len(w) < 3 or w in STOPWORDS:
            continue
        freq[w] = freq.get(w, 0) + 1

    bigrams = {}
    for i in range(len(words) - 1):
        w1, w2 = words[i], words[i + 1]
        if w1 in STOPWORDS or w2 in STOPWORDS or len(w1) < 3 or len(w2) < 3:
            continue
        phrase = f"{w1} {w2}"
        bigrams[phrase] = bigrams.get(phrase, 0) + 1

    ranked_words = sorted(freq.items(), key=lambda x: (-x[1], x[0]))
    ranked_bigrams = sorted(
        ((p, c) for p, c in bigrams.items() if c >= 2),
        key=lambda x: (-x[1], x[0]),
    )

    keywords, seen = [], set()
    for phrase, _ in ranked_bigrams[: top_n // 2]:
        if phrase not in seen:
            seen.add(phrase)
            keywords.append(phrase)
    for word, _ in ranked_words:
        if len(keywords) >= top_n:
            break
        if word not in seen and not any(word in k.split() for k in keywords):
            seen.add(word)
            keywords.append(word)

    return keywords[:top_n]


def flatten_profile_text(profile: dict) -> str:
    """Collapse the tailored profile into one lowercase text blob for matching."""
    parts = [
        profile.get("summary", ""),
        profile.get("headline", ""),
    ]
    for s in profile.get("skills", []) or []:
        parts.append(s.get("category", ""))
        parts.append(s.get("items", ""))
    for job in profile.get("experience", []) or []:
        parts.append(job.get("title", ""))
        for b in job.get("bullets", []) or []:
            parts.append(b)
    for e in profile.get("education", []) or []:
        parts.append(e)
    return " ".join(parts).lower()


def calculate_keyword_coverage(jd_text: str, tailored_profile: dict, top_n: int = 25):
    """
    Returns (coverage_percent, matched_list, missing_list) computed
    entirely in code from the JD and the FINAL tailored resume text —
    no dependence on the AI's self-reported number.
    """
    keywords = extract_jd_keywords(jd_text, top_n=top_n)
    resume_blob = flatten_profile_text(tailored_profile)

    matched, missing = [], []
    for kw in keywords:
        if kw in resume_blob:
            matched.append(kw)
        else:
            missing.append(kw)

    pct = round(100 * len(matched) / len(keywords)) if keywords else 0
    return pct, matched, missing


# ---------------------------------------------------------------
# 5. DOCX GENERATION  (clean, single-column, 100% ATS-parseable)
# ---------------------------------------------------------------
def build_docx(profile: dict, out_path: Path):
    doc = Document()

    section = doc.sections[0]
    section.page_width, section.page_height = Inches(8.5), Inches(11)
    for side in ("left", "right", "top", "bottom"):
        setattr(section, f"{side}_margin", Inches(0.5))

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(2)

    def para(align=None, before=0, after=2):
        p = doc.add_paragraph()
        if align:
            p.alignment = align
        p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after = Pt(after)
        return p

    def run(p, text, bold=False, italic=False, size=10):
        r = p.add_run(text)
        r.bold, r.italic = bold, italic
        r.font.size = Pt(size)
        r.font.name = "Times New Roman"
        return r

    def section_header(text):
        p = para(before=8, after=2)
        run(p, text.upper(), bold=True, size=10.5)

    p = para(WD_ALIGN_PARAGRAPH.CENTER)
    run(p, profile.get("name", ""), bold=True, size=14)
    if profile.get("headline"):
        p = para(WD_ALIGN_PARAGRAPH.CENTER)
        run(p, profile["headline"], bold=True, size=10.5)
    if profile.get("contact_line"):
        p = para(WD_ALIGN_PARAGRAPH.CENTER, after=4)
        run(p, profile["contact_line"])

    if profile.get("summary"):
        section_header("Professional Summary")
        p = para()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run(p, profile["summary"])

    if profile.get("skills"):
        section_header("Core Skills")
        for s in profile["skills"]:
            p = para()
            run(p, f"{s.get('category', '')}: ", bold=True)
            run(p, s.get("items", ""))

    if profile.get("experience"):
        section_header("Professional Experience")
        usable_width = section.page_width - section.left_margin - section.right_margin
        for job in profile["experience"]:
            p = para(before=6)
            p.paragraph_format.tab_stops.add_tab_stop(usable_width, WD_TAB_ALIGNMENT.RIGHT)
            run(p, f"{job.get('title', '')} — {job.get('company', '')}", bold=True)
            run(p, "\t")
            run(p, job.get("dates", ""), bold=True)
            for b in job.get("bullets", []):
                bp = doc.add_paragraph(style="List Bullet")
                bp.paragraph_format.space_after = Pt(1)
                bp.paragraph_format.left_indent = Inches(0.25)
                r = bp.add_run(b)
                r.font.size = Pt(10)
                r.font.name = "Times New Roman"

    if profile.get("education"):
        section_header("Education & Certifications")
        for line in profile["education"]:
            p = para()
            if "—" in line:
                left, right = line.split("—", 1)
                run(p, left.strip() + " ", bold=True)
                run(p, "— " + right.strip())
            else:
                run(p, line, bold=True)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def build_cover_letter_docx(profile: dict, cover_letter_text: str, company: str,
                             job_title: str, out_path: Path):
    """Simple, professional single-page cover letter matching the resume's font/margins."""
    doc = Document()

    section = doc.sections[0]
    section.page_width, section.page_height = Inches(8.5), Inches(11)
    for side in ("left", "right", "top", "bottom"):
        setattr(section, f"{side}_margin", Inches(1))

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(10)

    def para(align=None):
        p = doc.add_paragraph()
        if align:
            p.alignment = align
        return p

    def run(p, text, bold=False, size=11):
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(size)
        r.font.name = "Times New Roman"
        return r

    p = para()
    run(p, profile.get("name", ""), bold=True, size=13)
    if profile.get("contact_line"):
        p = para()
        run(p, profile["contact_line"], size=10)

    p = para()
    run(p, datetime.now().strftime("%B %d, %Y"))

    p = para()
    run(p, f"Re: Application for {job_title} at {company}", bold=True)

    p = para()
    run(p, "Dear Hiring Manager,")

    for block in (cover_letter_text or "").split("\n\n"):
        block = block.strip()
        if not block:
            continue
        p = para(WD_ALIGN_PARAGRAPH.JUSTIFY)
        run(p, block)

    p = para()
    run(p, "Sincerely,")
    p = para()
    run(p, profile.get("name", ""), bold=True)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def safe_filename(text: str) -> str:
    return re.sub(r"[^A-Za-z0-9]+", "_", text).strip("_")[:60]


def extract_forced_company(jd_block: str):
    """
    If a batch JD block starts with a line like "COMPANY: Acme Corp", pull that out
    as an override for the folder/company name and strip it from the JD text sent
    to the AI. Returns (forced_company_or_None, cleaned_jd_text).
    """
    lines = jd_block.splitlines()
    if lines and lines[0].strip().lower().startswith(FORCE_COMPANY_PREFIX):
        forced = lines[0].split(":", 1)[1].strip()
        remainder = "\n".join(lines[1:]).strip()
        return (forced or None), (remainder or jd_block)
    return None, jd_block


# ===============================================================
#                      STREAMLIT UI
# ===============================================================
st.title("⚡ ATS Resume Automator")
st.caption("Paste a JD → get a tailored, ATS-clean resume in seconds. "
           "Job titles, companies, dates & education are never changed. No fake facts. "
           "Multi-key Gemini pool + multi-provider automatic failover.")
st.markdown(
    f"""
    <div style="
        display:inline-flex; align-items:center; gap:8px;
        background:linear-gradient(135deg,#6366f1 0%,#8b5cf6 50%,#ec4899 100%);
        color:#ffffff; padding:6px 16px; border-radius:999px;
        font-size:13px; font-weight:600; margin:2px 0 14px 0;
        box-shadow:0 3px 10px rgba(139,92,246,0.35);
        letter-spacing:0.2px;">
        <span>👨‍💻</span>
        <span>Developed by {APP_AUTHOR}</span>
        <span style="opacity:0.8; font-weight:500;">· {APP_VERSION}</span>
    </div>
    """,
    unsafe_allow_html=True,
)

with st.sidebar:
    st.header("① One-time Setup")

    # ---------- Gemini multi-key pool ----------
    st.subheader("🔑 Gemini key pool (recommended)")
    st.caption("Upload a plain .txt file with ALL your Gemini API keys, one per line. "
               "The app rotates through them automatically when one gets rate-limited.")

    uploaded_keys_file = st.file_uploader(
        "Upload Gemini keys file (.txt)",
        type=["txt"],
        key="gemini_keys_uploader",
        help="One API key per line. Blank lines and lines starting with # are ignored.",
    )

    if uploaded_keys_file is not None:
        try:
            content = uploaded_keys_file.getvalue().decode("utf-8", errors="ignore")
            st.session_state["gemini_key_pool"] = parse_keys_text(content)
            st.session_state["gemini_key_pool_error"] = None
        except Exception as e:                        # noqa: BLE001
            st.session_state["gemini_key_pool"] = []
            st.session_state["gemini_key_pool_error"] = str(e)

    with st.expander("Advanced: load from a local file path instead"):
        st.caption("⚠️ This only works when you run the app on your own computer "
                   "(`streamlit run app.py`). It will NOT work on Streamlit Cloud or any "
                   "hosted deployment — the server has no access to your PC's C:\\ drive, "
                   "and pasting a web URL here will not work either since this reads local "
                   "files, not web pages. Use the uploader above instead on Streamlit Cloud.")
        gemini_key_file_path = st.text_input(
            "Gemini keys file (.txt, one key per line)",
            value="",
            placeholder=r"e.g. C:\Users\YourName\gemini_keys.txt",
            key="gemini_key_file_path",
            help="Type or paste the FULL path to your own keys file on this computer — "
                 "everyone's path will be different, so nothing is pre-filled here.",
        )
        reload_clicked = st.button("🔄 Reload keys from local path", use_container_width=True)
        if reload_clicked:
            try:
                loaded = load_keys_from_file(gemini_key_file_path) if gemini_key_file_path.strip() else []
                st.session_state["gemini_key_pool"] = loaded
                st.session_state["gemini_key_pool_error"] = None
            except Exception as e:                        # noqa: BLE001
                st.session_state["gemini_key_pool"] = []
                st.session_state["gemini_key_pool_error"] = str(e)

    gemini_key_pool = st.session_state.get("gemini_key_pool", [])
    pool_error = st.session_state.get("gemini_key_pool_error")

    if pool_error:
        st.error(f"Couldn't read key file: {pool_error}")
    elif gemini_key_pool:
        st.success(f"Loaded {len(gemini_key_pool)} Gemini key(s).")
    else:
        st.warning("No keys loaded yet — upload a .txt file above.")

    if gemini_key_pool:
        with st.expander(f"Show loaded keys (masked) — {len(gemini_key_pool)} total"):
            for i, k in enumerate(gemini_key_pool, start=1):
                st.markdown(f"{i}. `{gemini_key_short(k)}`")

    st.markdown("---")

    # ---------- Optional manual keys (single key each) ----------
    with st.expander("Optional: manual keys (fallback, single key each)"):
        st.caption("Only needed if you don't want to use the key-file pool above, "
                   "or want extra providers as backup.")
        keys = {}
        for p in PROVIDERS:
            keys[p["id"]] = st.text_input(
                p["label"], type="password", help=p["hint"], key=f"key_{p['id']}"
            )

    st.markdown("---")

    # ---------- Candidate path picker ----------
    st.subheader("👤 Candidate")

    st.markdown(
        "Upload the candidate's resume below — this works both locally and on "
        "Streamlit Cloud."
    )
    uploaded_resume_file = st.file_uploader(
        "📤 Upload resume",
        type=["docx", "pdf"],
        key="resume_uploader",
    )

    force_reparse = st.checkbox("Re-parse resume (use if you edited the original)", value=False)

    # ---------- Cooldown status ----------
    cooldowns = st.session_state.get("cooldowns", {})
    active_cd = {k: t for k, t in cooldowns.items() if t > time.time()}
    if active_cd:
        st.markdown("---")
        st.markdown("**⏳ Rate-limited (auto-retry soon):**")
        label_by_id = {p["id"]: p["label"] for p in PROVIDERS}
        for key, t in active_cd.items():
            mins = int((t - time.time()) // 60) + 1
            if key.startswith("gemini:"):
                raw_key = key.split("gemini:", 1)[1]
                st.markdown(f"- Gemini key `{gemini_key_short(raw_key)}` — {mins} min")
            else:
                st.markdown(f"- {label_by_id.get(key, key)} — {mins} min")

    # ---------- Footer credit ----------
    st.markdown("---")
    st.markdown(
        f"""
        <div style="
            text-align:center; padding:14px 10px; margin-top:4px;
            border-radius:14px;
            background:linear-gradient(135deg,rgba(99,102,241,0.15),rgba(236,72,153,0.15));
            border:1px solid rgba(139,92,246,0.35);">
            <div style="font-size:20px; margin-bottom:2px;">⚡</div>
            <div style="font-size:11px; letter-spacing:1px; opacity:0.7; text-transform:uppercase;">
                ATS Resume Automator {APP_VERSION}
            </div>
            <div style="
                font-size:15px; font-weight:800; margin-top:4px;
                background:linear-gradient(135deg,#6366f1,#8b5cf6,#ec4899);
                -webkit-background-clip:text; -webkit-text-fill-color:transparent;
                background-clip:text;">
                Developed by {APP_AUTHOR}
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

st.header("② Job Description(s)")

include_cover_letter = st.checkbox(
    "✉️ Also generate a tailored cover letter with every resume",
    value=True,
    help="Adds one extra field to the SAME AI call used to tailor the resume — "
         "this does not use a second API call, so it's free in terms of quota.",
)

tab_single, tab_batch = st.tabs(["📄 Single JD", "📦 Batch Mode (multiple JDs)"])

with tab_single:
    jd_text = st.text_area(
        "Job Description", height=280,
        placeholder="Copy the entire JD from LinkedIn / Indeed / Dice and paste here…",
        key="single_jd_text",
    )
    generate = st.button("🚀 Generate Tailored Resume", type="primary",
                         use_container_width=True, key="single_generate_btn")

with tab_batch:
    st.caption(
        "Process many job descriptions in one run. You'll get back a single ZIP with "
        "one folder per company (resume + cover letter if enabled above), plus a "
        "`summary.csv` listing coverage % and status for every application."
    )
    batch_input_mode = st.radio(
        "How do you want to provide the JDs?",
        ["Paste multiple JDs", "Upload JD .txt files"],
        horizontal=True,
        key="batch_input_mode",
    )

    batch_jd_items = []  # list of (label, raw_jd_block)

    if batch_input_mode == "Paste multiple JDs":
        st.caption(
            f"Paste JDs one after another. Put a line containing exactly "
            f"`{JD_SPLIT_MARKER}` between each JD. Optionally start any JD with a line "
            f"like `COMPANY: Acme Corp` to force that folder/company name, useful when "
            f"a JD (e.g. from a recruiter) doesn't clearly state the hiring company."
        )
        batch_jd_text = st.text_area(
            "Multiple Job Descriptions",
            height=320,
            placeholder=(
                f"COMPANY: Acme Corp\n<paste JD 1 here>\n\n{JD_SPLIT_MARKER}\n\n"
                f"<paste JD 2 here>\n\n{JD_SPLIT_MARKER}\n\n<paste JD 3 here>"
            ),
            key="batch_jd_text",
        )
        if batch_jd_text.strip():
            for i, block in enumerate(re.split(re.escape(JD_SPLIT_MARKER), batch_jd_text), start=1):
                block = block.strip()
                if len(block) >= 100:
                    batch_jd_items.append((f"JD {i}", block))
    else:
        batch_files = st.file_uploader(
            "Upload one .txt file per JD",
            type=["txt"],
            accept_multiple_files=True,
            key="batch_jd_files",
        )
        if batch_files:
            for f in batch_files:
                text = f.getvalue().decode("utf-8", errors="ignore").strip()
                if len(text) >= 100:
                    batch_jd_items.append((Path(f.name).stem, text))

    if batch_jd_items:
        st.info(f"📋 {len(batch_jd_items)} job description(s) ready to process.")

    generate_batch = st.button(
        f"🚀 Generate {len(batch_jd_items)} Tailored Resume(s)" if batch_jd_items else "🚀 Generate Batch",
        type="primary", use_container_width=True,
        disabled=not batch_jd_items, key="batch_generate_btn",
    )

if generate:
    have_gemini_pool = bool(gemini_key_pool)
    have_manual_key = any((v or "").strip() for v in keys.values())
    if not have_gemini_pool and not have_manual_key:
        st.error("Please load a Gemini keys file OR enter at least one manual provider key "
                 "in the sidebar.")
        st.stop()

    if uploaded_resume_file is None:
        st.error("Please upload a resume file (.docx or .pdf) in the sidebar.")
        st.stop()

    # Save the uploaded bytes to a temp file so the rest of the pipeline
    # (which expects a real filesystem Path) works unchanged — this is
    # what makes the app work on Streamlit Cloud, not just locally.
    tmp_dir = Path(tempfile.gettempdir()) / "ats_resume_uploads"
    tmp_dir.mkdir(parents=True, exist_ok=True)
    upload_name = Path(uploaded_resume_file.name).name  # strip any path junk
    resume_path = tmp_dir / upload_name
    with open(resume_path, "wb") as f:
        f.write(uploaded_resume_file.getbuffer())
    st.caption(f"🔎 Using uploaded file: `{upload_name}`")

    if len(jd_text.strip()) < 100:
        st.error("That JD looks too short — paste the full job description.")
        st.stop()
    t0 = time.time()

    try:
        with st.status("Working…", expanded=True) as status:
            st.write("📄 Reading candidate profile…")
            profile, parse_provider = load_or_parse_profile(
                keys, gemini_key_pool, resume_path, force=force_reparse
            )
            if parse_provider:
                st.write(f"　↳ profile learned via {parse_provider} (one-time)")

            st.write("🧠 Tailoring resume to this JD (8-stage process)…")
            tailored, provider_used, notes = tailor_profile(
                keys, gemini_key_pool, profile, jd_text,
                include_cover_letter=include_cover_letter,
            )
            for n in notes:
                st.write(f"　↳ {n}")
            st.write(f"　↳ generated via **{provider_used}**")

            st.write("📊 Calculating real keyword coverage (code-level, not AI-reported)…")
            coverage_pct, matched_kws, missing_kws = calculate_keyword_coverage(jd_text, tailored)
            st.write(f"　↳ {coverage_pct}% of top JD keywords found in the tailored resume")

            st.write("📝 Building ATS-compliant Word document…")
            # Keep the original resume filename exactly (always .docx output),
            # overwriting any previous tailored version for this candidate+company.
            fname = resume_path.stem + ".docx"

            company_for_jd = tailored.get("company_detected", "Company") or "Company"
            company_folder = safe_filename(company_for_jd) or "Company"

            out_dir = Path("output") / company_folder
            out_path = out_dir / fname

            build_docx(tailored, out_path)

            cover_letter_text = (tailored.get("cover_letter") or "").strip() if include_cover_letter else ""
            cl_path = None
            if cover_letter_text:
                st.write("✉️ Building tailored cover letter…")
                cl_path = out_dir / (resume_path.stem + "_CoverLetter.docx")
                build_cover_letter_docx(
                    profile, cover_letter_text, company_for_jd,
                    tailored.get("job_title_detected", "Role"), cl_path,
                )

            status.update(label=f"Done in {time.time()-t0:.1f}s ✅ (via {provider_used})",
                          state="complete")

        job_title = tailored.get("job_title_detected", "Role")
        company = tailored.get("company_detected", "Company")

        # ---- Package as a zip so the downloaded item unpacks into a
        # ---- folder named after the company (browsers can't save a
        # ---- single download "into" a folder directly — zip is the
        # ---- standard workaround for that). ----
        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
            zf.write(out_path, arcname=f"{company_folder}/{fname}")
            if cl_path and cl_path.exists():
                zf.write(cl_path, arcname=f"{company_folder}/{cl_path.name}")
        zip_bytes = zip_buffer.getvalue()
        zip_filename = f"{company_folder}.zip"

        contents_note = "the resume and cover letter" if cl_path else "the resume"
        st.success(
            f"**Generated:** `{out_path}` — packaged into **{zip_filename}**, which "
            f"unzips into a **{company_folder}/** folder containing {contents_note}.  \n"
            f"📁 Company folder: **{company_folder}**"
        )

        # ---- Attempt an automatic download (no click needed). Most
        # ---- browsers allow this since it's triggered right after a
        # ---- direct button click (Generate), but some may still block
        # ---- it once per site the first time — the manual button
        # ---- below always works as a fallback either way. ----
        b64_zip = base64.b64encode(zip_bytes).decode()
        auto_dl_id = f"autodl_{int(time.time() * 1000)}"
        st.components.v1.html(
            f"""
            <a id="{auto_dl_id}" href="data:application/zip;base64,{b64_zip}"
               download="{zip_filename}" style="display:none;"></a>
            <script>
                document.getElementById("{auto_dl_id}").click();
            </script>
            """,
            height=0,
            width=0,
        )

        col1, col2, col3 = st.columns([2, 1, 2])
        with col1:
            st.subheader("🎯 Detected Role")
            st.write(f"**{job_title}** at **{company}**")
            st.download_button(
                f"⬇️ Download {zip_filename} (folder: {company_folder}/)",
                zip_bytes,
                file_name=zip_filename,
                mime="application/zip",
                use_container_width=True,
            )
            st.caption("If the download didn't start automatically, use the button above. "
                       "Unzip it and you'll get a folder named after the company with the "
                       "resume inside.")
        with col2:
            st.subheader("📊 Coverage")
            st.metric("JD keywords (code-calculated)", f"{coverage_pct}%",
                      help="Calculated in code from actual JD word frequency vs. the final "
                           "resume text — not self-reported by the AI. Target 70–80%; "
                           "100% looks like keyword stuffing to recruiters.")
        with col3:
            st.subheader("🔑 Keywords Matched")
            st.write(" · ".join(matched_kws) if matched_kws else "—")

        with st.expander("📋 Full keyword coverage detail (code-calculated)", expanded=False):
            c1, c2 = st.columns(2)
            with c1:
                st.markdown(f"**✅ Matched ({len(matched_kws)})**")
                for kw in matched_kws or ["—"]:
                    st.markdown(f"- {kw}")
            with c2:
                st.markdown(f"**❌ Missing ({len(missing_kws)})**")
                for kw in missing_kws or ["—"]:
                    st.markdown(f"- {kw}")

        gap = tailored.get("gap_analysis", {}) or {}
        with st.expander("🔍 Gap Analysis (Stage 3) — how the JD was matched", expanded=True):
            g1, g2, g3 = st.columns(3)
            with g1:
                st.markdown("**✅ Direct matches**")
                for item in gap.get("direct_matches", []) or ["—"]:
                    st.markdown(f"- {item}")
            with g2:
                st.markdown("**🔁 Implied matches (made explicit)**")
                for item in gap.get("implied_matches", []) or ["—"]:
                    st.markdown(f"- {item}")
            with g3:
                st.markdown("**⛔ Not added (candidate lacks this)**")
                for item in (gap.get("no_match_not_fabricated", [])
                             or tailored.get("missing_keywords", []) or ["—"]):
                    st.markdown(f"- {item}")
            st.caption("⛔ items were deliberately NOT claimed — never state these in the "
                       "application or interview prep either.")

        with st.expander("Preview new summary & skills"):
            st.markdown(f"**Summary:** {tailored.get('summary','')}")
            for s in tailored.get("skills", []):
                st.markdown(f"**{s.get('category','')}:** {s.get('items','')}")

        if cover_letter_text:
            with st.expander("✉️ Preview cover letter", expanded=False):
                for block in cover_letter_text.split("\n\n"):
                    if block.strip():
                        st.markdown(block.strip())
                        st.markdown("")

    except Exception as e:                            # noqa: BLE001
        st.error(f"Something went wrong: {e}")
        st.info("Common fixes: add more keys to your Gemini key file, check your API keys, "
                "add another provider key in the sidebar, check internet, or tick "
                "'Re-parse resume' if the profile cache is corrupted.")

if generate_batch:
    have_gemini_pool = bool(gemini_key_pool)
    have_manual_key = any((v or "").strip() for v in keys.values())
    if not have_gemini_pool and not have_manual_key:
        st.error("Please load a Gemini keys file OR enter at least one manual provider key "
                 "in the sidebar.")
        st.stop()

    if uploaded_resume_file is None:
        st.error("Please upload a resume file (.docx or .pdf) in the sidebar.")
        st.stop()

    tmp_dir = Path(tempfile.gettempdir()) / "ats_resume_uploads"
    tmp_dir.mkdir(parents=True, exist_ok=True)
    upload_name = Path(uploaded_resume_file.name).name
    resume_path = tmp_dir / upload_name
    with open(resume_path, "wb") as f:
        f.write(uploaded_resume_file.getbuffer())
    st.caption(f"🔎 Using uploaded file: `{upload_name}` — {len(batch_jd_items)} job description(s)")

    batch_t0 = time.time()

    try:
        with st.spinner("📄 Reading candidate profile…"):
            profile, parse_provider = load_or_parse_profile(
                keys, gemini_key_pool, resume_path, force=force_reparse
            )
        if parse_provider:
            st.caption(f"　↳ profile learned via {parse_provider} (one-time)")
    except Exception as e:                            # noqa: BLE001
        st.error(f"Couldn't read/parse the resume: {e}")
        st.stop()

    results = []
    used_folder_names = set()
    progress = st.progress(0.0, text=f"Processing 0/{len(batch_jd_items)}…")
    log_box = st.empty()
    log_lines = []

    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        for idx, (label, raw_block) in enumerate(batch_jd_items, start=1):
            forced_company, jd_clean = extract_forced_company(raw_block)
            row = {"label": label, "company": "", "job_title": "", "coverage": None,
                  "status": "❌ failed", "provider": "", "error": ""}
            try:
                if len(jd_clean.strip()) < 100:
                    raise ValueError("JD text too short after removing the COMPANY: line")

                tailored, provider_used, notes = tailor_profile(
                    keys, gemini_key_pool, profile, jd_clean,
                    include_cover_letter=include_cover_letter,
                )
                coverage_pct, matched_kws, missing_kws = calculate_keyword_coverage(jd_clean, tailored)

                company_for_jd = forced_company or tailored.get("company_detected", "Company") or "Company"
                job_title = tailored.get("job_title_detected", "Role")

                company_folder = safe_filename(company_for_jd) or f"Company_{idx}"
                n = 2
                while company_folder in used_folder_names:
                    company_folder = f"{safe_filename(company_for_jd) or 'Company'}_{n}"
                    n += 1
                used_folder_names.add(company_folder)

                fname = resume_path.stem + ".docx"
                out_dir = Path("output") / "batch" / company_folder
                out_path = out_dir / fname
                build_docx(tailored, out_path)
                zf.write(out_path, arcname=f"{company_folder}/{fname}")

                cl_text = (tailored.get("cover_letter") or "").strip() if include_cover_letter else ""
                if cl_text:
                    cl_path = out_dir / (resume_path.stem + "_CoverLetter.docx")
                    build_cover_letter_docx(profile, cl_text, company_for_jd, job_title, cl_path)
                    zf.write(cl_path, arcname=f"{company_folder}/{cl_path.name}")

                row.update({
                    "company": company_for_jd, "job_title": job_title,
                    "coverage": coverage_pct, "status": "✅ done",
                    "provider": provider_used,
                })
                log_lines.append(f"✅ {label} → {company_for_jd} / {job_title} — "
                                 f"{coverage_pct}% via {provider_used}")
            except Exception as e:                    # noqa: BLE001
                row["error"] = str(e)[:200]
                log_lines.append(f"❌ {label} — failed: {row['error'][:150]}")

            results.append(row)
            log_box.code("\n".join(log_lines[-12:]) or "…", language=None)
            progress.progress(idx / len(batch_jd_items),
                              text=f"Processed {idx}/{len(batch_jd_items)} — {label}")

        summary_io = io.StringIO()
        writer = csv.writer(summary_io)
        writer.writerow(["Label", "Company", "Job Title", "Coverage %", "Status", "Provider", "Error"])
        for r in results:
            writer.writerow([r["label"], r["company"], r["job_title"],
                             r.get("coverage", ""), r["status"], r.get("provider", ""), r.get("error", "")])
        zf.writestr("summary.csv", summary_io.getvalue())

    zip_bytes = zip_buffer.getvalue()
    batch_zip_name = f"BatchApplications_{datetime.now().strftime('%Y%m%d_%H%M')}.zip"
    ok_count = sum(1 for r in results if r["status"].startswith("✅"))

    if ok_count == len(results):
        st.success(f"Finished all {len(results)} job description(s) in "
                  f"{time.time()-batch_t0:.1f}s ✅")
    else:
        st.warning(f"Finished {len(results)} job description(s) in {time.time()-batch_t0:.1f}s — "
                  f"{ok_count} succeeded, {len(results)-ok_count} failed (see table below).")

    st.download_button(
        f"⬇️ Download {batch_zip_name} ({ok_count} companies + summary.csv)",
        zip_bytes, file_name=batch_zip_name, mime="application/zip",
        use_container_width=True,
    )

    st.subheader("📋 Batch summary")
    st.dataframe(
        [{"Company": r["company"] or "—", "Role": r["job_title"] or "—",
          "Coverage %": r.get("coverage") if r.get("coverage") is not None else "—",
          "Status": r["status"], "Provider": r.get("provider") or "—",
          "Notes": r.get("error", "")}
         for r in results],
        use_container_width=True,
    )

st.markdown("---")
st.markdown(
    f"""
    <div style="
        text-align:center; padding:22px 10px 8px 10px;">
        <div style="font-size:13px; opacity:0.65;">
            ⚡ <strong>ATS Resume Automator</strong> &nbsp;{APP_VERSION}
        </div>
        <div style="font-size:17px; margin-top:6px; font-weight:500;">
            Crafted with <span style="color:#ec4899;">♥</span> by
            <span style="
                font-weight:800; font-size:19px; margin-left:3px;
                background:linear-gradient(135deg,#6366f1,#8b5cf6,#ec4899);
                -webkit-background-clip:text; -webkit-text-fill-color:transparent;
                background-clip:text;">
                {APP_AUTHOR}
            </span>
        </div>
    </div>
    """,
    unsafe_allow_html=True,
)
